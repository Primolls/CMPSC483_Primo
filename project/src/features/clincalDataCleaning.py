import os
import re
import pandas as pd
import numpy as np
from docx import Document
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestRegressor



# change the current working directory from 'project\src\features' to 'project\volume\data\raw'
os.chdir(os.path.abspath('..'))
os.chdir(os.path.abspath('..\\volume\\data\\raw'))

# Load clinical data from MS word
part1 = Document('Scores 1-39.docx')
part2 = Document('Scores 40-74.docx')

# change the current working directory from 'project\volume\data\raw' to 'project\src\features'
os.chdir(os.path.abspath('..'))
os.chdir(os.path.abspath('..'))
os.chdir(os.path.abspath('..\\src\\features'))

treatments0 = []
for para in part1.paragraphs:
    treatments0.append(para.text)

# part 2 was not labeled
# for para in part2.paragraphs:
#     treatments.append(para.text)

# remove empty strings
treatments0 = [x for x in treatments0 if x != '']

treatments = []
patientId = []

# extract Id and corresponding treatments
for i in range(len(treatments0)):
    tempId = treatments0[i].split(')', 1)[0][1:].strip()
    tempTreatments = treatments0[i].split(':', 1)[1].strip()#.split(';')
    # for j in range(len(tempTreatments)):
    #     tempTreatments[j] = tempTreatments[j].strip()
    treatments.append(tempTreatments)
    patientId.append(tempId)


# print all cells in all tables
# for table in part1.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             print(cell.text)


# visualize contents and their location
# for i in range(len(part1.tables)):
#     for j in range(len(part1.tables[i].rows)):
#         for k in range(len(part1.tables[i].rows[j].cells)):
#             print('table:', i, 'row:', j, 'cell', k, part1.tables[i].rows[j].cells[k].text)

# row 17 cell 3-7

# create a dataframe to store part 1
clinicalDf1 = pd.DataFrame(columns = ['childPugh', 'okuda', 'bclc', 'hklc', 'clip'])

for i in range(len(part1.tables)):
    tempList = []
    for j in range(3,8):
        tempList.append(part1.tables[i].rows[17].cells[j].text)
    clinicalDf1.loc[i] = tempList


# create a dataframe to store part 2
clinicalDf2 = pd.DataFrame(columns = ['childPugh', 'okuda', 'bclc', 'hklc', 'clip'])

for i in range(len(part2.tables)):
    tempList = []
    for j in range(3,8):
        tempList.append(part2.tables[i].rows[17].cells[j].text)
    clinicalDf2.loc[i] = tempList

# combine two parts
# but we only use the df for part 1, therefore we comment out instead of delete this line
# clinicalDf = clinicalDf1.append(clinicalDf2)

# add treatments and patientId, which are stored in lists
clinicalDf1['treatments'] = treatments
# clinicalDf1['patientId'] = patientId

# save features with labels
features = clinicalDf1

# see how the updated df looks
# print(clinicalDf1['treatments'])


############################################################################
# the following code refers to:                                       #
# https://towardsdatascience.com/random-forest-in-python-24d0893d51c0      #
############################################################################


# deal with dummy variables
features = pd.get_dummies(features)
# store labels
label1 = np.array(features['treatments_BCLC- LT or PEI; HKLC- Resection/LT/Ablation'])
label2 = np.array(features['treatments_BCLC- No treatment; HKLC- Supportive Care'])
label3 = np.array(features['treatments_BCLC- Resection; HKLC- Resection/LT/Ablation'])
label4 = np.array(features['treatments_BCLC- Resection; HKLC- TACE'])
label5 = np.array(features['treatments_BCLC- TACE; HKLC- LT'])
label6 = np.array(features['treatments_BCLC- TACE; HKLC- Resection/LT/Ablation'])
label7 = np.array(features['treatments_BCLC- TACE; HKLC- TACE'])
label8 = np.array(features['treatments_BCLC-TACE; HKLC-Resection/LT/Ablation'])

# save label names
labelNames = features.columns[-8:]

# remove labels from  features
features= features.drop(columns=features.columns[-8:])

# save treatment names
feature_list = list(features.columns)

# convert to numpy array
features = np.array(features)

# use sklearn library for further analyze
# including split training and testing sets
# and model training



def TrainModel(features, labels, testSize = 0.25, randomState = 25, nEstimators = 100):
    print('----------------Process Starts----------------')
    train_features, test_features, train_labels, test_labels = train_test_split(features, labels, test_size = testSize, random_state = randomState)

    # print('Training Features Shape:', train_features.shape)
    # print('Training Labels Shape:', train_labels.shape)
    # print('Testing Features Shape:', test_features.shape)
    # print('Testing Labels Shape:', test_labels.shape)

    # outputs for above codes
    # Training Features Shape: (26, 26)
    # Training Labels Shape: (26,)
    # Testing Features Shape: (9, 26)
    # Testing Labels Shape: (9,)


    # Train Model
    print('----Step: Train Model Starts----')
    rf = RandomForestRegressor(n_estimators = nEstimators, random_state = randomState)
    rf.fit(train_features, train_labels)
    print('----Step: Train Model Ends----')

    # Make Predictions on the Test Set
    print('----Step: Make Prediction Starts----')
    predictions = rf.predict(test_features)
    errors = abs(predictions - test_labels)
    print('Mean Absolute Error:', round(np.mean(errors), 2), 'degrees.')
    print('----Step: Make Prediction Ends----')


    # Determine Performance Metrics
    print('----Step: Determine Performance Starts----')
    mape = 100 * (errors / test_labels)
    accuracy = 100 - np.mean(mape)
    print('Accuracy:', round(accuracy, 2), '%.')
    print('----Step: Determine Performance Ends----')

    # step: return the model for saving purpose should be added in future developments

print(labelNames)
print(feature_list)

# each line predicts a single treatment in the following list (labelNames):
#        'treatments_BCLC- LT or PEI; HKLC- Resection/LT/Ablation',
#        'treatments_BCLC- No treatment; HKLC- Supportive Care',
#        'treatments_BCLC- Resection; HKLC- Resection/LT/Ablation',
#        'treatments_BCLC- Resection; HKLC- TACE',
#        'treatments_BCLC- TACE; HKLC- LT',
#        'treatments_BCLC- TACE; HKLC- Resection/LT/Ablation',
#        'treatments_BCLC- TACE; HKLC- TACE',
#        'treatments_BCLC-TACE; HKLC-Resection/LT/Ablation'


# TrainModel(features, label1)
# TrainModel(features, label2)
# TrainModel(features, label3)
# TrainModel(features, label4)
# TrainModel(features, label5)
# TrainModel(features, label6)
# TrainModel(features, label7)
# TrainModel(features, label8)


# Issues:
# 1) in line 165, the value of test_labels is zero, which leads to error
#       I might did something wrong and caused this issue
# 2) There are few data but too many different treatments to predict
# Possible solutions:
# 1. treat BCLC and HKLC separately
#           which requires extra working on dataframe: treatments0
#           lines commented out in 41-43 should be useful for this purpose
# 2. label patients 40-74, and add data patients 1, 2, 3, 20 to the dataset
# 3. There might be some other solutions